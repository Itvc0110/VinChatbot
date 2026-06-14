from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Any
from urllib.parse import urljoin, urlparse

from bs4 import BeautifulSoup

from vinchatbot.app.core.config import get_settings
from vinchatbot.app.ingest.assets import (
    build_binary_image_asset_record,
    build_file_asset_record,
    build_ocr_text_record,
    extract_html_image_assets,
    extract_html_table_records,
    extract_pdf_page_image_assets,
    file_asset_records_from_links,
    is_image_url,
    parse_csv_records,
    table_record_to_text,
)
from vinchatbot.app.ingest.normalizer import (
    classify_domain,
    infer_category,
    infer_source_kind,
    normalize_text,
)
from vinchatbot.app.ingest.ocr import OcrResult, ocr_dependency_status, run_english_ocr_image
from vinchatbot.app.schemas.document import (
    CalendarEvent,
    LinkReference,
    RawDocument,
    SourceDocumentMetadata,
    StructuredRecord,
    stable_hash,
)

logger = logging.getLogger(__name__)

PARSER_VERSION = "v3"
MONTHS = {
    "jan": 1,
    "feb": 2,
    "mar": 3,
    "apr": 4,
    "may": 5,
    "jun": 6,
    "jul": 7,
    "aug": 8,
    "sep": 9,
    "sept": 9,
    "oct": 10,
    "nov": 11,
    "dec": 12,
}
DATE_TEXT_PATTERN = re.compile(
    r"(?P<date>"
    r"\b\d{1,2}(?:-\d{1,2})?[- ](?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)"
    r"(?:[- ]\d{2,4})?"
    r"(?:-\d{1,2}[- ](?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)(?:[- ]\d{2,4})?)?"
    r"|\b\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?"
    r"|\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\s+\d{4}"
    r")",
    flags=re.IGNORECASE,
)
DATE_VALUE_PATTERN = re.compile(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)", re.I)
VND_PATTERN = re.compile(r"(?P<amount>\d[\d,.\s]*)\s*(?P<currency>VND|VNĐ)", re.IGNORECASE)


def parse_html(
    content: str,
    source_url: str,
    source_metadata: SourceDocumentMetadata | None = None,
) -> RawDocument:
    settings = get_settings()
    title = source_url
    soup = BeautifulSoup(content, "html.parser")
    if soup.title and soup.title.text:
        title = normalize_text(soup.title.text)

    source_kind = source_metadata.source_kind if source_metadata else infer_source_kind(source_url, title=title)
    extracted = _extract_html_text(content, soup, source_kind)
    metadata = _base_raw_metadata(source_url, title, source_kind, source_metadata)
    metadata.update(extract_policy_detail_metadata(content, source_url))
    metadata["link_count"] = len(extract_links_from_html(content, source_url))

    structured_records = extract_structured_records_from_html(
        content,
        source_url,
        metadata,
        enable_image_asset_extraction=settings.enable_image_asset_extraction,
        enable_ocr=settings.enable_ocr,
        image_download_enabled=settings.image_download_enabled,
        ocr_engine=settings.ocr_engine,
        ocr_model=settings.ocr_model,
        ocr_lang=settings.ocr_lang,
    )

    if source_metadata:
        source_metadata.parser_name = "html"
        source_metadata.parser_version = PARSER_VERSION
        source_metadata.content_hash = stable_hash(extracted or "")

    return RawDocument(
        source_url=source_url,
        canonical_url=metadata.get("canonical_url", source_url),
        title=metadata.get("document_title") or title,
        document_type=source_kind if source_kind != "external_public_page" else "html",
        content=normalize_text(extracted or ""),
        metadata=metadata,
        source_metadata=source_metadata,
        structured_records=structured_records,
    )


def _extract_html_text(content: str, soup: BeautifulSoup, source_kind: str) -> str:
    if source_kind in {
        "gateway_page",
        "policy_listing",
        "policy_html",
        "financial_policy",
        "academic_catalog",
        "calendar_page",
        "registrar_page",
        "library_page",
    }:
        return _extract_structured_html_text(soup)

    markdown = get_settings().enable_markdown_parsing
    extracted = None
    try:
        import trafilatura

        kwargs = {"include_links": False, "include_tables": True}
        if markdown:
            kwargs["output_format"] = "markdown"
        extracted = trafilatura.extract(content, **kwargs)
    except (ImportError, TypeError):
        try:
            import trafilatura

            extracted = trafilatura.extract(content, include_links=False, include_tables=True)
        except ImportError:
            extracted = None

    if not extracted:
        extracted = _extract_structured_html_text(soup)
    return extracted


def _extract_structured_html_text(soup: BeautifulSoup) -> str:
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()
    headings_and_text: list[str] = []
    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "p", "li", "td", "th"]):
        text = normalize_text(element.get_text(" ", strip=True))
        if not text:
            continue
        if element.name in {"h1", "h2", "h3", "h4", "h5"}:
            level = int(element.name[1])
            headings_and_text.append(f"{'#' * level} {text}")
        elif element.name == "li":
            headings_and_text.append(f"- {text}")
        else:
            headings_and_text.append(text)
    return "\n\n".join(headings_and_text)


def _pdf_pages_to_markdown(document) -> dict[int, str] | None:
    """Per-page Markdown via pymupdf4llm (preserves headings/lists/tables). Returns None on
    failure so the caller falls back to plain text."""
    try:
        import pymupdf4llm

        chunks = pymupdf4llm.to_markdown(document, page_chunks=True, show_progress=False)
        return {index + 1: (chunk.get("text") or "") for index, chunk in enumerate(chunks)}
    except Exception:
        logger.debug("pymupdf4llm markdown extraction failed; using plain text.", exc_info=True)
        return None


def _docx_to_markdown(stream) -> str:
    import docx

    document = docx.Document(stream)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "") or ""
        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
            except (ValueError, IndexError):
                level = 2
            lines.append(f"{'#' * max(1, min(6, level))} {text}")
        elif style.startswith("List") or text.lstrip().startswith(("-", "•", "*")):
            lines.append(f"- {text.lstrip('-•* ').strip()}")
        else:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(lines)


def parse_docx(
    content: bytes,
    source_url: str,
    source_metadata: SourceDocumentMetadata | None = None,
) -> RawDocument:
    try:
        import docx  # noqa: F401
    except ImportError as exc:
        raise RuntimeError("Install python-docx to parse DOCX files.") from exc

    title = Path(urlparse(source_url).path).name or source_url
    source_kind = source_metadata.source_kind if source_metadata else infer_source_kind(source_url, title=title)
    markdown = _docx_to_markdown(io.BytesIO(content))
    metadata = _base_raw_metadata(source_url, title, source_kind, source_metadata)
    metadata["parser_name"] = "docx"
    metadata["file_size_bytes"] = len(content)
    metadata["content_hash"] = stable_hash(content.hex())
    if source_metadata:
        source_metadata.parser_name = "docx"
        source_metadata.parser_version = PARSER_VERSION
        source_metadata.content_hash = metadata.get("content_hash")
        source_metadata.file_size_bytes = len(content)
    return RawDocument(
        source_url=source_url,
        canonical_url=metadata.get("canonical_url", source_url),
        title=normalize_text(title),
        document_type="markdown",
        content=normalize_text(markdown),
        metadata=metadata,
        source_metadata=source_metadata,
        structured_records=[],
    )


def parse_pdf_bytes(
    content: bytes,
    source_url: str,
    title: str | None = None,
    source_metadata: SourceDocumentMetadata | None = None,
) -> RawDocument:
    settings = get_settings()
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("Install pymupdf to parse PDF files.") from exc

    pages: list[str] = []
    page_stats: list[dict[str, Any]] = []
    ocr_results: dict[int, OcrResult] = {}
    ocr_status_by_page: dict[int, str] = {}
    with fitz.open(stream=content, filetype="pdf") as document:
        meta_title = document.metadata.get("title") if document.metadata else None
        title = title or meta_title or Path(source_url).name or source_url
        page_markdown = _pdf_pages_to_markdown(document) if settings.enable_pdf_markdown else None
        for index, page in enumerate(document, start=1):
            markdown = page_markdown.get(index) if page_markdown else None
            page_text = normalize_text(markdown if markdown else page.get_text("text"))
            text_char_count = len(page_text)
            image_count = len(page.get_images(full=True))
            page_stats.append(
                {
                    "page_number": index,
                    "text_char_count": text_char_count,
                    "embedded_image_count": image_count,
                    "width": round(page.rect.width),
                    "height": round(page.rect.height),
                }
            )
            if page_text:
                pages.append(f"# Trang {index}\n{page_text}")
            if (
                settings.enable_ocr
                and index <= settings.ocr_max_pdf_pages
                and text_char_count < settings.ocr_min_text_chars_per_page
            ):
                available, reason = ocr_dependency_status(settings.ocr_engine)
                if not available:
                    ocr_status_by_page[index] = reason or "skipped_dependency_missing"
                    continue
                try:
                    pixmap = page.get_pixmap(dpi=200, alpha=False)
                    result = run_english_ocr_image(
                        pixmap.tobytes("png"),
                        lang=settings.ocr_lang,
                        model=settings.ocr_model,
                        store_boxes=settings.ocr_store_boxes,
                    )
                    if result.text:
                        ocr_results[index] = result
                        ocr_status_by_page[index] = "completed"
                    else:
                        ocr_status_by_page[index] = "skipped_no_text"
                except Exception as exc:  # pragma: no cover - depends on optional OCR runtime
                    logger.warning("OCR failed for %s page %s: %s", source_url, index, exc)
                    ocr_status_by_page[index] = "failed"

    source_kind = source_metadata.source_kind if source_metadata else infer_source_kind(source_url, "application/pdf", title)
    extracted = normalize_text("\n\n".join(pages))
    metadata = _base_raw_metadata(source_url, title, source_kind, source_metadata)
    metadata["parser_name"] = "pdf"
    metadata["file_size_bytes"] = len(content)
    metadata["content_hash"] = stable_hash(content.hex())
    metadata["page_count"] = len(pages)
    if source_kind == "calendar_pdf":
        metadata.setdefault("academic_year", infer_academic_year(extracted, source_url))
    metadata["needs_ocr"] = any(
        int(stat.get("text_char_count") or 0) < settings.ocr_min_text_chars_per_page
        for stat in page_stats
    )

    raw = RawDocument(
        source_url=source_url,
        canonical_url=metadata.get("canonical_url", source_url),
        title=normalize_text(title),
        document_type=source_kind if source_kind != "external_public_page" else "pdf",
        content=extracted,
        metadata=metadata,
        source_metadata=source_metadata,
    )
    raw.structured_records = structured_records_from_raw(raw)
    raw.structured_records.extend(
        extract_pdf_page_image_assets(
            page_stats=page_stats,
            source_url=source_url,
            parent_doc_id=raw.parent_doc_id,
            metadata=metadata,
            enable_ocr=settings.enable_ocr,
            ocr_engine=settings.ocr_engine,
            ocr_model=settings.ocr_model,
            ocr_lang=settings.ocr_lang,
            ocr_min_text_chars_per_page=settings.ocr_min_text_chars_per_page,
            ocr_max_pdf_pages=settings.ocr_max_pdf_pages,
            ocr_results=ocr_results,
            ocr_status_by_page=ocr_status_by_page,
        )
    )
    raw.structured_records.extend(extract_pdf_table_records(content, source_url, raw.parent_doc_id, metadata))
    if source_metadata:
        source_metadata.parser_name = "pdf"
        source_metadata.parser_version = PARSER_VERSION
        source_metadata.content_hash = raw.content_hash
        source_metadata.file_size_bytes = len(content)
    return raw


def parse_binary_asset_bytes(
    content: bytes,
    source_url: str,
    content_type: str | None = None,
    source_metadata: SourceDocumentMetadata | None = None,
) -> RawDocument:
    settings = get_settings()
    title = Path(urlparse(source_url).path).name or source_url
    source_kind = source_metadata.source_kind if source_metadata else infer_source_kind(source_url, content_type, title)
    metadata = _base_raw_metadata(source_url, title, source_kind, source_metadata)
    metadata["file_size_bytes"] = len(content)
    metadata["mime_type"] = content_type.split(";", 1)[0] if content_type else None
    metadata["content_hash"] = stable_hash(content.hex())
    parent_doc_id = stable_hash(metadata.get("canonical_url") or source_url)
    structured_records: list[StructuredRecord] = []
    ocr_result: OcrResult | None = None
    ocr_status: str | None = None

    if is_image_url(source_url, metadata.get("mime_type")):
        if settings.enable_ocr and len(content) <= settings.ocr_max_image_mb * 1024 * 1024:
            available, reason = ocr_dependency_status(settings.ocr_engine)
            if available:
                try:
                    ocr_result = run_english_ocr_image(
                        content,
                        lang=settings.ocr_lang,
                        model=settings.ocr_model,
                        store_boxes=settings.ocr_store_boxes,
                    )
                    ocr_status = "completed" if ocr_result.text else "skipped_no_text"
                except Exception as exc:  # pragma: no cover - depends on optional OCR runtime
                    logger.warning("OCR failed for image %s: %s", source_url, exc)
                    ocr_status = "failed"
            else:
                ocr_status = reason or "skipped_dependency_missing"
        elif settings.enable_ocr:
            ocr_status = "skipped_size_cap"
        image_record = build_binary_image_asset_record(
            source_url=source_url,
            parent_doc_id=parent_doc_id,
            metadata=metadata,
            content=content,
            mime_type=metadata.get("mime_type"),
            enable_ocr=settings.enable_ocr,
            ocr_engine=settings.ocr_engine,
            ocr_model=settings.ocr_model,
            ocr_lang=settings.ocr_lang,
            ocr_result=ocr_result,
            ocr_status=ocr_status,
        )
        structured_records.append(image_record)
        if ocr_result and ocr_result.text:
            structured_records.append(
                build_ocr_text_record(
                    source_url=source_url,
                    parent_doc_id=parent_doc_id,
                    metadata=metadata,
                    asset_url=source_url,
                    page_number=None,
                    result=ocr_result,
                    ocr_engine=settings.ocr_engine,
                    ocr_model=settings.ocr_model,
                    ocr_lang=settings.ocr_lang,
                )
            )
    else:
        structured_records.append(
            build_file_asset_record(
                source_url=source_url,
                parent_doc_id=parent_doc_id,
                metadata=metadata,
                content=content,
                mime_type=metadata.get("mime_type"),
            )
        )

    raw = RawDocument(
        source_url=source_url,
        canonical_url=metadata.get("canonical_url", source_url),
        title=normalize_text(title),
        document_type=source_kind if source_kind in {"image_asset", "file_asset"} else "file_asset",
        content=ocr_result.text if ocr_result and ocr_result.text else "",
        metadata=metadata,
        source_metadata=source_metadata,
        structured_records=structured_records,
    )
    if source_metadata:
        source_metadata.parser_name = "binary_asset"
        source_metadata.parser_version = PARSER_VERSION
        source_metadata.content_hash = stable_hash(content.hex())
        source_metadata.file_size_bytes = len(content)
    return raw


def parse_markdown(
    content: str,
    source_url: str,
    source_metadata: SourceDocumentMetadata | None = None,
) -> RawDocument:
    title = _markdown_title(content) or Path(urlparse(source_url).path).name or source_url
    source_kind = source_metadata.source_kind if source_metadata else "markdown"
    extracted = normalize_text(content)
    metadata = _base_raw_metadata(source_url, title, source_kind, source_metadata)
    metadata["parser_name"] = "markdown"
    structured_records = extract_markdown_table_records(extracted, source_url, metadata)
    raw = RawDocument(
        source_url=source_url,
        canonical_url=metadata.get("canonical_url", source_url),
        title=normalize_text(title),
        document_type="markdown",
        content=extracted,
        metadata=metadata,
        source_metadata=source_metadata,
        structured_records=structured_records,
    )
    if source_metadata:
        source_metadata.parser_name = "markdown"
        source_metadata.parser_version = PARSER_VERSION
        source_metadata.content_hash = raw.content_hash
    return raw


def parse_spreadsheet_bytes(
    content: bytes,
    source_url: str,
    content_type: str | None = None,
    source_metadata: SourceDocumentMetadata | None = None,
) -> RawDocument:
    title = Path(urlparse(source_url).path).name or source_url
    source_kind = source_metadata.source_kind if source_metadata else infer_source_kind(source_url, content_type, title)
    metadata = _base_raw_metadata(source_url, title, source_kind, source_metadata)
    metadata["parser_name"] = "spreadsheet"
    metadata["file_size_bytes"] = len(content)
    metadata["mime_type"] = content_type.split(";", 1)[0] if content_type else None
    metadata["content_hash"] = stable_hash(content.hex())
    extension = Path(urlparse(source_url).path).suffix.lower()

    if _looks_like_html_spreadsheet(content):
        metadata["parser_warning"] = "spreadsheet_html_fallback"
        html_text = content.decode("utf-8-sig", errors="replace")
        raw = parse_html(html_text, source_url, source_metadata=source_metadata)
        raw.metadata.update(metadata)
        raw.metadata["parser_name"] = "spreadsheet_html_fallback"
        raw.document_type = "spreadsheet"
        if source_metadata:
            source_metadata.parser_name = "spreadsheet_html_fallback"
            source_metadata.parser_version = PARSER_VERSION
            source_metadata.content_hash = metadata.get("content_hash") or raw.content_hash
            source_metadata.file_size_bytes = len(content)
        return raw

    if (
        extension == ".csv"
        or (content_type or "").lower().startswith("text/csv")
        or _looks_like_csv_spreadsheet(content)
    ):
        if extension != ".csv" and not (content_type or "").lower().startswith("text/csv"):
            metadata["parser_warning"] = "spreadsheet_csv_fallback"
        extracted, records = parse_csv_records(content, source_url, metadata)
    else:
        extracted, records = _parse_excel_records(content, source_url, metadata, extension=extension)
    records.extend(extract_domain_records_from_spreadsheet(records, metadata))

    raw = RawDocument(
        source_url=source_url,
        canonical_url=metadata.get("canonical_url", source_url),
        title=normalize_text(title),
        document_type="csv" if extension == ".csv" else "spreadsheet",
        content=extracted,
        metadata=metadata,
        source_metadata=source_metadata,
        structured_records=records,
    )
    if source_metadata:
        source_metadata.parser_name = "spreadsheet"
        source_metadata.parser_version = PARSER_VERSION
        source_metadata.content_hash = metadata.get("content_hash") or raw.content_hash
        source_metadata.file_size_bytes = len(content)
    return raw


DATE_PATTERN = DATE_TEXT_PATTERN


def extract_calendar_events(raw: RawDocument) -> list[CalendarEvent]:
    events: list[CalendarEvent] = []
    academic_year = raw.metadata.get("academic_year")
    term_hint = raw.metadata.get("term")
    current_page: int | None = None

    for line in raw.content.splitlines():
        stripped = normalize_text(line)
        if not stripped:
            continue
        page_match = re.match(r"#\s*Trang\s+(\d+)", stripped, flags=re.IGNORECASE)
        if page_match:
            current_page = int(page_match.group(1))
            continue
        match = DATE_PATTERN.search(stripped)
        if not match:
            continue
        lowered = stripped.lower()
        if not any(
            keyword in lowered
            for keyword in [
                "instruction",
                "deadline",
                "drop",
                "add",
                "exam",
                "holiday",
                "spring",
                "fall",
                "summer",
                "học kỳ",
                "kỳ thi",
                "hạn",
            ]
        ):
            continue
        term = term_hint
        if "fall" in lowered:
            term = "Fall"
        elif "spring" in lowered:
            term = "Spring"
        elif "summer" in lowered:
            term = "Summer"
        date_text = match.group("date")
        date_start_iso, date_end_iso = normalize_date_range(date_text, raw.metadata.get("academic_year"))
        events.append(
            CalendarEvent(
                event_name=stripped,
                date_start=date_text,
                event_type=infer_event_type(stripped),
                date_text_original=date_text,
                date_start_iso=date_start_iso,
                date_end_iso=date_end_iso,
                academic_year=academic_year,
                term=term,
                source_url=raw.source_url,
                page_number=current_page,
            )
        )
    return events


def _base_raw_metadata(
    source_url: str,
    title: str,
    source_kind: str,
    source_metadata: SourceDocumentMetadata | None = None,
) -> dict[str, Any]:
    domain, domain_type, source_trust = classify_domain(source_url)
    category, subcategory = infer_category(source_url, title)
    if source_kind == "policy_listing":
        category, subcategory = "policy", "listing"
    elif source_kind == "gateway_page":
        category, subcategory = "gateway", "student_support"
    elif source_kind in {"calendar_page", "calendar_pdf"}:
        category, subcategory = "academic", "calendar"
    elif source_kind == "financial_policy":
        category, subcategory = "student_affairs", "financial"
    elif source_kind == "academic_catalog":
        category, subcategory = "academic", "catalog"
    elif source_kind == "registrar_page":
        category, subcategory = "academic", "registrar"
    elif source_kind == "library_page":
        category, subcategory = "student_services", "library"
    elif source_kind in {"image_asset", "file_asset", "spreadsheet", "csv", "markdown"}:
        category, subcategory = "asset", source_kind

    metadata = {
        "source_kind": source_kind,
        "source_id": stable_hash(source_url),
        "canonical_url": source_url,
        "document_title": normalize_text(title),
        "domain": domain,
        "domain_type": domain_type,
        "source_trust": source_trust,
        "category": category,
        "subcategory": subcategory,
        "parser_name": "html",
        "parser_version": PARSER_VERSION,
    }
    if source_metadata:
        metadata.update(source_metadata.model_dump(exclude_none=True))
        metadata["source_id"] = source_metadata.source_id
        metadata["canonical_url"] = source_metadata.canonical_url
    return metadata


def extract_links_from_html(content: str, source_url: str) -> list[LinkReference]:
    soup = BeautifulSoup(content, "html.parser")
    links: list[LinkReference] = []
    section_path: list[str] = []
    for element in soup.find_all(["h1", "h2", "h3", "h4", "a"]):
        if element.name in {"h1", "h2", "h3", "h4"}:
            level = int(element.name[1])
            heading = normalize_text(element.get_text(" ", strip=True))
            if heading:
                section_path = section_path[: level - 1]
                section_path.append(heading)
            continue
        href = element.get("href")
        if not href:
            continue
        target_url = urljoin(source_url, href)
        if target_url.startswith("mailto:") or target_url.startswith("tel:"):
            continue
        anchor_text = normalize_text(element.get_text(" ", strip=True))
        context = _nearby_text(element)
        domain, domain_type, source_trust = classify_domain(target_url)
        links.append(
            LinkReference(
                source_url=source_url,
                target_url=target_url.split("#", 1)[0],
                anchor_text=anchor_text or None,
                link_context=context or None,
                section_path=list(section_path),
                discovered_from=source_url,
                domain=domain,
                domain_type=domain_type,
                source_trust=source_trust,
            )
        )
    return links


def parse_policy_listing_records(content: str, source_url: str, metadata: dict[str, Any] | None = None) -> list[StructuredRecord]:
    metadata = metadata or {}
    soup = BeautifulSoup(content, "html.parser")
    records: list[StructuredRecord] = []
    for link in soup.find_all("a", href=True):
        href = str(link["href"])
        if "/all-policies/" not in href and "/publication/" not in href:
            continue
        title = normalize_text(link.get_text(" ", strip=True))
        if not title or title.lower() in {"all policies", "publication / public"}:
            continue
        target_url = urljoin(source_url, href)
        row_text = normalize_text(link.find_parent("tr").get_text(" ", strip=True)) if link.find_parent("tr") else _nearby_text(link)
        dates = re.findall(r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{2},\s+\d{4}", row_text)
        policy_code = _extract_policy_code(row_text, title, dates)
        data = {
            "policy_title": title,
            "policy_code": policy_code,
            "date_issued": dates[0] if dates else None,
            "date_last_updated": dates[1] if len(dates) > 1 else (dates[0] if dates else None),
            "detail_url": target_url,
            "listing_category": metadata.get("subcategory") or _listing_category_from_url(source_url),
        }
        record_id = stable_hash(f"policy_listing:{source_url}:{target_url}:{policy_code or title}")
        records.append(
            StructuredRecord(
                record_id=record_id,
                record_type="policy_listing",
                parent_doc_id=stable_hash(source_url),
                source_url=source_url,
                title=title,
                data=data,
                metadata=metadata,
            )
        )
    return _dedupe_records(records)


def extract_policy_detail_metadata(content: str, source_url: str) -> dict[str, Any]:
    soup = BeautifulSoup(content, "html.parser")
    text_lines = [normalize_text(line) for line in soup.get_text("\n", strip=True).splitlines()]
    text_lines = [line for line in text_lines if line]
    result: dict[str, Any] = {}
    h1 = soup.find("h1")
    if h1:
        result["document_title"] = normalize_text(h1.get_text(" ", strip=True))

    label_map = {
        "Reference Number": "reference_number",
        "Document Type": "document_status",
        "Issuing By": "issuing_unit",
        "Issuing Date": "issued_date",
        "Date Issued": "issued_date",
        "Date Last Updated": "updated_date",
        "Applying for": "applying_for",
        "Security Classification": "security_classification",
    }
    for index, line in enumerate(text_lines):
        clean_label = line.rstrip(":")
        if clean_label not in label_map:
            continue
        value = _next_value(text_lines, index + 1)
        if value:
            result[label_map[clean_label]] = value

    pdf_link = None
    for link in soup.find_all("a", href=True):
        href = urljoin(source_url, str(link["href"]))
        if href.lower().endswith(".pdf") or "wp-content/uploads" in href:
            pdf_link = href
    if pdf_link:
        result["pdf_url"] = pdf_link
    return result


def extract_structured_records_from_html(
    content: str,
    source_url: str,
    metadata: dict[str, Any] | None = None,
    *,
    enable_image_asset_extraction: bool = True,
    enable_ocr: bool = False,
    image_download_enabled: bool = False,
    ocr_engine: str = "paddleocr",
    ocr_model: str = "PP-OCRv5",
    ocr_lang: str = "en",
) -> list[StructuredRecord]:
    metadata = metadata or {}
    source_kind = metadata.get("source_kind") or infer_source_kind(source_url)
    records: list[StructuredRecord] = []
    if source_kind == "policy_listing":
        records.extend(parse_policy_listing_records(content, source_url, metadata))
    records.extend(extract_html_table_records(content, source_url, metadata))
    if enable_image_asset_extraction:
        records.extend(
            extract_html_image_assets(
                content,
                source_url,
                metadata,
                enable_ocr=enable_ocr,
                image_download_enabled=image_download_enabled,
                ocr_engine=ocr_engine,
                ocr_model=ocr_model,
                ocr_lang=ocr_lang,
            )
        )

    raw = RawDocument(
        source_url=source_url,
        canonical_url=metadata.get("canonical_url", source_url),
        title=metadata.get("document_title") or source_url,
        document_type=source_kind if source_kind != "external_public_page" else "html",
        content=_extract_structured_html_text(BeautifulSoup(content, "html.parser")),
        metadata=metadata,
    )
    records.extend(structured_records_from_raw(raw))
    if source_kind in {"gateway_page", "registrar_page", "library_page", "external_public_page"}:
        links = extract_links_from_html(content, source_url)
        records.extend(link_records_from_links(links, raw.parent_doc_id, metadata))
        records.extend(file_asset_records_from_links(links, raw.parent_doc_id, metadata))
    return _dedupe_records(records)


def structured_records_from_raw(raw: RawDocument) -> list[StructuredRecord]:
    records: list[StructuredRecord] = []
    if raw.metadata.get("source_kind") in {"calendar_page", "calendar_pdf"} or "calendar" in raw.title.lower():
        for event in extract_calendar_events(raw):
            records.append(
                StructuredRecord(
                    record_id=stable_hash(f"calendar_event:{raw.parent_doc_id}:{event.event_name}"),
                    record_type="calendar_event",
                    parent_doc_id=raw.parent_doc_id,
                    source_url=raw.source_url,
                    title=event.event_name,
                    data=event.model_dump(),
                    metadata=raw.metadata,
                )
            )
    if raw.metadata.get("source_kind") == "financial_policy" or "financial" in raw.title.lower():
        records.extend(extract_fee_records(raw))
    if raw.metadata.get("source_kind") == "academic_catalog":
        records.extend(extract_program_records(raw))
    return _dedupe_records(records)


def extract_fee_records(raw: RawDocument) -> list[StructuredRecord]:
    records: list[StructuredRecord] = []
    for index, line in enumerate(raw.content.splitlines()):
        stripped = normalize_text(line)
        if not stripped:
            continue
        amount_match = VND_PATTERN.search(stripped)
        if not amount_match:
            continue
        amount_text = amount_match.group("amount").strip()
        currency = amount_match.group("currency").upper().replace("VNĐ", "VND")
        fee_name = stripped[: amount_match.start()].strip(" :-") or stripped
        data = {
            "fee_name": fee_name,
            "fee_type": infer_fee_type(stripped),
            "amount": normalize_amount(amount_text),
            "currency": currency,
            "amount_text_original": amount_match.group(0),
            "collection_time": infer_collection_time(stripped),
            "applies_to": None,
            "conditions": stripped,
            "row_index": index,
        }
        records.append(
            StructuredRecord(
                record_id=stable_hash(f"fee_record:{raw.parent_doc_id}:{index}:{stripped}"),
                record_type="fee_record",
                parent_doc_id=raw.parent_doc_id,
                source_url=raw.source_url,
                title=fee_name,
                data=data,
                metadata=raw.metadata,
            )
        )
    return records


def extract_program_records(raw: RawDocument) -> list[StructuredRecord]:
    records: list[StructuredRecord] = []
    current_college: str | None = None
    current_degree: str | None = None
    for line in raw.content.splitlines():
        stripped = normalize_text(line).lstrip("#").strip()
        if not stripped:
            continue
        if stripped.lower().startswith("college of"):
            current_college = stripped
        if any(prefix in stripped for prefix in ["Bachelor of", "Master of", "Ph.D.", "Medical Doctor"]):
            current_degree = "undergraduate" if "Bachelor" in stripped or "Medical Doctor" in stripped else "graduate"
            data = {
                "college": current_college,
                "degree_level": current_degree,
                "program_name": stripped,
                "major": stripped,
                "concentration": None,
                "minor": None,
                "cohort": infer_cohort(stripped),
            }
            records.append(
                StructuredRecord(
                    record_id=stable_hash(f"program:{raw.parent_doc_id}:{current_college}:{stripped}"),
                    record_type="program",
                    parent_doc_id=raw.parent_doc_id,
                    source_url=raw.source_url,
                    title=stripped,
                    data=data,
                    metadata=raw.metadata,
                )
            )
    return records


def extract_pdf_table_records(
    content: bytes,
    source_url: str,
    parent_doc_id: str,
    metadata: dict[str, Any],
) -> list[StructuredRecord]:
    try:
        import pdfplumber
    except ImportError:
        return []

    records: list[StructuredRecord] = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as document:
            for page_index, page in enumerate(document.pages, start=1):
                tables = page.extract_tables() or []
                for table_index, table in enumerate(tables, start=1):
                    rows = [
                        [normalize_text(str(cell or "")) for cell in row]
                        for row in table
                        if row and any(str(cell or "").strip() for cell in row)
                    ]
                    if not rows:
                        continue
                    headers = rows[0]
                    data_rows = rows[1:] if len(rows) > 1 else []
                    table_id = stable_hash(f"pdf_table:{source_url}:{page_index}:{table_index}")
                    rag_text = table_record_to_text(
                        title=f"{metadata.get('document_title') or source_url} page {page_index}",
                        headers=headers,
                        rows=data_rows,
                    )
                    record_metadata = dict(metadata)
                    record_metadata.update({"page_number": page_index, "table_id": table_id})
                    records.append(
                        StructuredRecord(
                            record_id=stable_hash(f"table_record:{table_id}"),
                            record_type="table_record",
                            parent_doc_id=parent_doc_id,
                            source_url=source_url,
                            title=f"PDF table page {page_index}",
                            data={
                                "table_id": table_id,
                                "caption": None,
                                "headers": headers,
                                "rows": data_rows,
                                "row_count": len(data_rows),
                                "column_count": len(headers),
                                "section_path": list(metadata.get("section_path", [])),
                                "page_number": page_index,
                                "rag_text": rag_text,
                                "content_hash": stable_hash(rag_text),
                            },
                            metadata=record_metadata,
                        )
                    )
    except Exception as exc:  # pragma: no cover - depends on optional parser/runtime
        logger.warning("pdfplumber table extraction failed for %s: %s", source_url, exc)
    return records


def extract_markdown_table_records(
    content: str,
    source_url: str,
    metadata: dict[str, Any],
) -> list[StructuredRecord]:
    parent_doc_id = stable_hash(metadata.get("canonical_url") or source_url)
    records: list[StructuredRecord] = []
    section_path: list[str] = []
    pending_table: list[list[str]] = []
    table_index = 0

    def flush() -> None:
        nonlocal pending_table, table_index
        if len(pending_table) < 2:
            pending_table = []
            return
        headers = pending_table[0]
        data_rows = pending_table[2:] if _is_markdown_separator(pending_table[1]) else pending_table[1:]
        if not data_rows:
            pending_table = []
            return
        table_index += 1
        table_id = stable_hash(f"markdown_table:{source_url}:{table_index}:{headers}")
        rag_text = table_record_to_text(
            title=metadata.get("document_title") or source_url,
            headers=headers,
            rows=data_rows,
        )
        records.append(
            StructuredRecord(
                record_id=stable_hash(f"table_record:{table_id}"),
                record_type="table_record",
                parent_doc_id=parent_doc_id,
                source_url=source_url,
                title=f"Markdown table {table_index}",
                data={
                    "table_id": table_id,
                    "caption": None,
                    "headers": headers,
                    "rows": data_rows,
                    "row_count": len(data_rows),
                    "column_count": len(headers),
                    "section_path": list(section_path),
                    "page_number": None,
                    "rag_text": rag_text,
                    "content_hash": stable_hash(rag_text),
                },
                metadata=metadata,
            )
        )
        pending_table = []

    for raw_line in content.splitlines():
        line = raw_line.strip()
        if line.startswith("#"):
            flush()
            level = max(1, min(6, len(line) - len(line.lstrip("#"))))
            heading = normalize_text(line.lstrip("#").strip())
            section_path = section_path[: level - 1]
            if heading:
                section_path.append(heading)
            continue
        if "|" in line and line.strip("|").strip():
            pending_table.append([normalize_text(cell) for cell in line.strip("|").split("|")])
            continue
        flush()
    flush()
    return records


def extract_domain_records_from_spreadsheet(
    spreadsheet_rows: list[StructuredRecord],
    metadata: dict[str, Any],
) -> list[StructuredRecord]:
    records: list[StructuredRecord] = []
    for row_record in spreadsheet_rows:
        if row_record.record_type != "spreadsheet_row":
            continue
        row_text = str(row_record.data.get("rag_text") or "")
        lowered = row_text.lower()
        if VND_PATTERN.search(row_text):
            raw = RawDocument(
                source_url=row_record.source_url,
                canonical_url=metadata.get("canonical_url", row_record.source_url),
                title=metadata.get("document_title") or row_record.title or row_record.source_url,
                document_type="spreadsheet",
                content=row_text,
                metadata={**metadata, "source_kind": "financial_policy"},
            )
            records.extend(extract_fee_records(raw))
            continue
        if DATE_PATTERN.search(row_text) and any(
            keyword in lowered
            for keyword in ["deadline", "exam", "holiday", "registration", "fall", "spring", "summer"]
        ):
            date_match = DATE_PATTERN.search(row_text)
            if not date_match:
                continue
            date_text = date_match.group("date")
            date_start_iso, date_end_iso = normalize_date_range(date_text, metadata.get("academic_year"))
            data = {
                "event_name": row_text,
                "event_type": infer_event_type(row_text),
                "date_text_original": date_text,
                "date_start_iso": date_start_iso,
                "date_end_iso": date_end_iso,
                "term": _term_from_text(row_text),
                "academic_year": metadata.get("academic_year"),
            }
            records.append(
                StructuredRecord(
                    record_id=stable_hash(f"calendar_event:{row_record.record_id}:{date_text}"),
                    record_type="calendar_event",
                    parent_doc_id=row_record.parent_doc_id,
                    source_url=row_record.source_url,
                    title=row_text[:120],
                    data=data,
                    metadata=metadata,
                )
            )
    return records


def _parse_excel_records(
    content: bytes,
    source_url: str,
    metadata: dict[str, Any],
    extension: str | None = None,
) -> tuple[str, list[StructuredRecord]]:
    try:
        import pandas as pd
    except ImportError:
        parent_doc_id = stable_hash(metadata.get("canonical_url") or source_url)
        return "", [
            build_file_asset_record(
                source_url=source_url,
                parent_doc_id=parent_doc_id,
                metadata={**metadata, "parser_warning": "spreadsheet_dependency_missing"},
                content=content,
                mime_type=metadata.get("mime_type"),
            )
        ]

    engine = _excel_engine_for_content(content, extension or Path(urlparse(source_url).path).suffix.lower())
    try:
        sheets = pd.read_excel(io.BytesIO(content), sheet_name=None, dtype=str, engine=engine)
    except Exception as exc:
        parent_doc_id = stable_hash(metadata.get("canonical_url") or source_url)
        warning = f"spreadsheet_parse_failed:{exc.__class__.__name__}"
        logger.warning("Spreadsheet parse failed for %s: %s", source_url, exc)
        return "", [
            build_file_asset_record(
                source_url=source_url,
                parent_doc_id=parent_doc_id,
                metadata={**metadata, "parser_warning": warning},
                content=content,
                mime_type=metadata.get("mime_type"),
            )
        ]
    parent_doc_id = stable_hash(metadata.get("canonical_url") or source_url)
    records: list[StructuredRecord] = []
    lines: list[str] = [f"Spreadsheet: {metadata.get('document_title') or source_url}"]
    for sheet_name, frame in sheets.items():
        frame = frame.fillna("")
        headers = [normalize_text(str(column)) for column in frame.columns]
        lines.append(f"Sheet: {sheet_name}")
        lines.append(f"Columns: {' | '.join(headers)}")
        for row_index, row in enumerate(frame.astype(str).values.tolist(), start=1):
            values = [normalize_text(value) for value in row]
            row_data = {
                header or f"column_{index + 1}": values[index] if index < len(values) else ""
                for index, header in enumerate(headers)
            }
            row_text = f"Sheet {sheet_name} row {row_index}: " + " | ".join(
                f"{key}: {value}" for key, value in row_data.items() if value
            )
            lines.append(row_text)
            records.append(
                StructuredRecord(
                    record_id=stable_hash(f"spreadsheet_row:{source_url}:{sheet_name}:{row_index}:{row_text}"),
                    record_type="spreadsheet_row",
                    parent_doc_id=parent_doc_id,
                    source_url=source_url,
                    title=f"{sheet_name} row {row_index}",
                    data={
                        "sheet_name": str(sheet_name),
                        "row_index": row_index,
                        "headers": headers,
                        "values": row_data,
                        "rag_text": row_text,
                        "content_hash": stable_hash(row_text),
                    },
                    metadata=metadata,
                )
            )
    return normalize_text("\n".join(lines)), records


def _looks_like_html_spreadsheet(content: bytes) -> bool:
    sample = content[:4096].lstrip().lower()
    return sample.startswith((b"<!doctype html", b"<html", b"<table")) or b"<table" in sample


def _looks_like_csv_spreadsheet(content: bytes) -> bool:
    if content.startswith((b"PK\x03\x04", b"\xd0\xcf\x11\xe0")) or b"\x00" in content[:4096]:
        return False
    sample = content[:4096].decode("utf-8-sig", errors="replace")
    if "<html" in sample.lower() or "<table" in sample.lower():
        return False
    if "\n" not in sample and "\r" not in sample:
        return False
    first_line = sample.splitlines()[0] if sample.splitlines() else ""
    return any(delimiter in first_line for delimiter in [",", "\t", ";"])


def _excel_engine_for_content(content: bytes, extension: str | None) -> str | None:
    extension = (extension or "").lower()
    if extension == ".xlsx" or content.startswith(b"PK\x03\x04"):
        return "openpyxl"
    if extension == ".xls" or content.startswith(b"\xd0\xcf\x11\xe0"):
        return "xlrd"
    return None


def link_records_from_links(
    links: list[LinkReference],
    parent_doc_id: str,
    metadata: dict[str, Any] | None = None,
) -> list[StructuredRecord]:
    metadata = metadata or {}
    records: list[StructuredRecord] = []
    for link in links:
        records.append(
            StructuredRecord(
                record_id=stable_hash(f"link_reference:{link.source_url}:{link.target_url}:{link.anchor_text}"),
                record_type="link_reference",
                parent_doc_id=parent_doc_id,
                source_url=link.source_url,
                title=link.anchor_text,
                data=link.model_dump(),
                metadata=metadata,
            )
        )
    return records


def normalize_date_range(date_text: str, academic_year: str | None = None) -> tuple[str | None, str | None]:
    normalized = date_text.replace("Sept", "Sep")
    parts = normalized.split("-")
    if not DATE_VALUE_PATTERN.search(normalized):
        return None, None

    years = _academic_year_bounds(academic_year)
    if len(parts) >= 3 and parts[1].isalpha():
        start = _date_token_to_iso(f"{parts[0]}-{parts[1]}", years)
        end = None
        if len(parts) >= 4:
            end = _date_token_to_iso(f"{parts[2]}-{parts[3]}", years)
        return start, end
    if len(parts) >= 2 and DATE_VALUE_PATTERN.search(parts[-1]):
        if len(parts) == 2:
            return _date_token_to_iso(normalized, years), None
        month = parts[-1]
        start = _date_token_to_iso(f"{parts[0]}-{month}", years)
        end = _date_token_to_iso(f"{parts[1]}-{month}", years)
        return start, end
    return _date_token_to_iso(normalized, years), None


def _date_token_to_iso(token: str, academic_year_bounds: tuple[int, int] | None) -> str | None:
    token = token.strip().replace(" ", "-")
    match = re.search(r"(?P<day>\d{1,2})[- ](?P<month>[A-Za-z]+)(?:[- ](?P<year>\d{2,4}))?", token)
    if not match:
        return None
    month = MONTHS.get(match.group("month").lower())
    if not month:
        return None
    year_text = match.group("year")
    if year_text:
        year = int(year_text)
        if year < 100:
            year += 2000
    elif academic_year_bounds:
        start_year, end_year = academic_year_bounds
        year = start_year if month >= 9 else end_year
    else:
        year = 2025 if month >= 9 else 2026
    return f"{year:04d}-{month:02d}-{int(match.group('day')):02d}"


def _academic_year_bounds(academic_year: str | None) -> tuple[int, int] | None:
    if not academic_year:
        return None
    years = re.findall(r"\d{4}", academic_year)
    if len(years) >= 2:
        return int(years[0]), int(years[1])
    return None


def infer_academic_year(text: str, url: str) -> str | None:
    match = re.search(r"(20\d{2})\s*[-/]\s*(20\d{2})", f"{text} {url}")
    if match:
        return f"{match.group(1)}-{match.group(2)}"
    if "2025" in text and "2026" in text:
        return "2025-2026"
    return None


def infer_event_type(text: str) -> str:
    lowered = text.lower()
    if "drop" in lowered:
        return "course_drop_deadline"
    if "add" in lowered or "transfer credit" in lowered:
        return "add_transfer_deadline"
    if "instruction begins" in lowered:
        return "instruction_begins"
    if "exam" in lowered:
        return "exam_period"
    if "holiday" in lowered or "tet" in lowered or "lunar" in lowered:
        return "holiday"
    if "registration" in lowered or "enrollment" in lowered:
        return "registration"
    return "academic_event"


def infer_fee_type(text: str) -> str:
    lowered = text.lower()
    if "tuition" in lowered:
        return "tuition"
    if "exam" in lowered:
        return "exam"
    if "library" in lowered or "document" in lowered or "equipment" in lowered:
        return "library"
    if "administrative" in lowered or "certified" in lowered or "student id" in lowered:
        return "academic_admin"
    return "fee"


def normalize_amount(value: str) -> int | None:
    digits = re.sub(r"\D", "", value)
    return int(digits) if digits else None


def infer_collection_time(text: str) -> str | None:
    lowered = text.lower()
    markers = ["at the time", "upon receipt", "within", "collected"]
    for marker in markers:
        index = lowered.find(marker)
        if index >= 0:
            return text[index:].strip()
    return None


def _markdown_title(content: str) -> str | None:
    for line in content.splitlines():
        stripped = normalize_text(line)
        if stripped.startswith("# "):
            return stripped.lstrip("#").strip()
    return None


def _is_markdown_separator(row: list[str]) -> bool:
    return all(set(cell.replace(":", "").strip()) <= {"-"} for cell in row if cell.strip())


def _term_from_text(text: str) -> str | None:
    lowered = text.lower()
    if "fall" in lowered:
        return "Fall"
    if "spring" in lowered:
        return "Spring"
    if "summer" in lowered:
        return "Summer"
    return None


def infer_cohort(text: str) -> str | None:
    match = re.search(r"(?:Cohort|AY)\s*([\d\s–-]+)", text, flags=re.IGNORECASE)
    return match.group(0) if match else None


def _extract_policy_code(row_text: str, title: str, dates: list[str]) -> str | None:
    text = row_text.replace(title, "", 1).strip()
    for date in dates:
        text = text.replace(date, " ")
    text = re.sub(r"\s+", " ", text).strip(" |:-")
    match = re.search(r"\b[A-Z0-9_./-]{2,}(?:-[A-Z0-9.]+)*\b", text)
    return match.group(0) if match else None


def _nearby_text(element) -> str:
    parent = element.find_parent(["tr", "li", "p", "div", "section"])
    if not parent:
        return normalize_text(element.get_text(" ", strip=True))
    return normalize_text(parent.get_text(" ", strip=True))


def _next_value(lines: list[str], start_index: int) -> str | None:
    for line in lines[start_index : start_index + 4]:
        if line and not line.endswith(":"):
            return line
    return None


def _listing_category_from_url(url: str) -> str:
    path = urlparse(url).path.strip("/")
    return path or "all-policies"


def _dedupe_records(records: list[StructuredRecord]) -> list[StructuredRecord]:
    seen: set[str] = set()
    deduped: list[StructuredRecord] = []
    for record in records:
        if record.record_id in seen:
            continue
        seen.add(record.record_id)
        deduped.append(record)
    return deduped
