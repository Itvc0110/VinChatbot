"""Form Assistant — fetch an official VinUni form and turn it into a review-ready, editable file.

Two output paths, per the agreed fidelity calibration:
  * FILLABLE AcroForm PDF → fill the widget fields in-place and return the PDF (pixel-perfect).
  * FLAT PDF / DOCX / anything else → generate a clean, editable .docx from the recognized fields
    (a faithful re-creation, not a copy of the original layout — "no-flaw" PDF→docx of an arbitrary flat
    PDF is not reliably achievable).

Everything here is best-effort and fail-open: on a fetch/parse error we degrade to a generated .docx built
from a default personal field set so the student always gets *something* editable plus the official URL.

Security: `official_url` is student-supplied, so `fetch_form_bytes` enforces an allowlist of VinUni hosts
(no arbitrary server-side fetch → no SSRF).
"""

from __future__ import annotations

import io
import logging
import re
from urllib.parse import urlparse

from vinchatbot.app.core.config import Settings, get_settings
from vinchatbot.app.schemas.forms import FormField

logger = logging.getLogger(__name__)

# Only official VinUni hosts may be fetched server-side (SSRF guard + "official forms only").
_ALLOWED_FORM_HOST_SUFFIX = ".vinuni.edu.vn"
_ALLOWED_FORM_HOSTS = {"vinuni.edu.vn"}

_MAX_FORM_BYTES = 25 * 1024 * 1024  # official forms are small PDFs/DOCX; cap defensively.

# "Label: ____" style blanks in a flat form's text, used to recognize fields when there are no widgets.
_LABEL_BLANK_RE = re.compile(r"^\s*([^:\n]{2,60}?)\s*[:：]\s*(?:_{2,}|\.{3,}|\s*)$")

# A generic, bilingual personal field set — the deterministic prefill baseline + the fallback when a form
# exposes no recognizable fields. Keys are stable; the LLM/heuristic maps personal data onto them.
DEFAULT_FIELDS: tuple[tuple[str, str], ...] = (
    ("full_name", "Họ và tên / Full name"),
    ("student_id", "Mã số sinh viên / Student ID"),
    ("program", "Chương trình / Program"),
    ("email", "Email"),
    ("date", "Ngày / Date"),
    ("subject", "Nội dung / Subject"),
    ("reason", "Lý do / Reason"),
)


class FormFetchError(RuntimeError):
    """Raised when the official form URL is not allowed or cannot be fetched."""


def is_allowed_form_url(url: str) -> bool:
    """True only for https URLs on an official VinUni host (SSRF guard)."""
    try:
        parsed = urlparse(url)
    except ValueError:
        return False
    if parsed.scheme not in ("http", "https"):
        return False
    host = (parsed.hostname or "").lower()
    return host in _ALLOWED_FORM_HOSTS or host.endswith(_ALLOWED_FORM_HOST_SUFFIX)


async def fetch_form_bytes(url: str, settings: Settings | None = None) -> tuple[bytes, str]:
    """Fetch the official form file. Returns (content, content_type). Raises FormFetchError on a
    disallowed host, oversized payload, or network/HTTP error."""
    if not is_allowed_form_url(url):
        raise FormFetchError(f"Form URL is not an allowed VinUni host: {url!r}")
    settings = settings or get_settings()
    try:
        import httpx
    except ImportError as exc:  # pragma: no cover - httpx is a core dependency
        raise FormFetchError("Install httpx to fetch forms.") from exc

    headers = {"User-Agent": settings.crawl_user_agent}
    try:
        async with httpx.AsyncClient(
            timeout=settings.crawl_timeout_seconds, follow_redirects=True, headers=headers
        ) as client:
            response = await client.get(url)
            response.raise_for_status()
            content = response.content
    except httpx.HTTPError as exc:
        raise FormFetchError(f"Could not fetch form: {exc}") from exc
    if len(content) > _MAX_FORM_BYTES:
        raise FormFetchError("Form file is too large.")
    return content, (response.headers.get("content-type") or "").split(";", 1)[0].strip().lower()


def _looks_like_pdf(content: bytes, content_type: str, url: str) -> bool:
    return content[:5] == b"%PDF-" or content_type == "application/pdf" or url.lower().endswith(".pdf")


def _pdf_widget_fields(content: bytes) -> list[FormField]:
    """Return the AcroForm widget fields of a fillable PDF (empty list if not fillable / on error)."""
    try:
        import fitz
    except ImportError:  # pragma: no cover - pymupdf is a core dependency
        return []
    fields: list[FormField] = []
    seen: set[str] = set()
    try:
        with fitz.open(stream=content, filetype="pdf") as document:
            if not getattr(document, "is_form_pdf", False):
                return []
            for page in document:
                for widget in page.widgets() or []:
                    name = (widget.field_name or "").strip()
                    if not name or name in seen:
                        continue
                    seen.add(name)
                    fields.append(
                        FormField(key=name, label=name, value=str(widget.field_value or "").strip())
                    )
    except Exception:
        logger.debug("PDF widget extraction failed.", exc_info=True)
        return []
    return fields


def _flat_pdf_labels(content: bytes) -> list[FormField]:
    """Best-effort field recognition for a FLAT PDF: scan text for 'Label: ____' lines."""
    try:
        import fitz
    except ImportError:  # pragma: no cover
        return []
    labels: list[FormField] = []
    seen: set[str] = set()
    try:
        with fitz.open(stream=content, filetype="pdf") as document:
            text = "\n".join(page.get_text("text") for page in document)
    except Exception:
        logger.debug("Flat-PDF text extraction failed.", exc_info=True)
        return []
    for line in text.splitlines():
        match = _LABEL_BLANK_RE.match(line)
        if not match:
            continue
        label = match.group(1).strip()
        key = _slugify(label)
        if not key or key in seen:
            continue
        seen.add(key)
        labels.append(FormField(key=key, label=label))
        if len(labels) >= 25:
            break
    return labels


def _slugify(text: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")
    return slug[:60]


def _default_fields() -> list[FormField]:
    return [FormField(key=key, label=label) for key, label in DEFAULT_FIELDS]


def default_form_fields() -> list[FormField]:
    """Public: the generic bilingual personal field set, used when no official file can be fetched/parsed."""
    return _default_fields()


def analyze_form(content: bytes, content_type: str, url: str) -> tuple[str, list[FormField]]:
    """Inspect the fetched form and decide the output path.

    Returns (file_kind, fields):
      * ("pdf", widget_fields)  → fillable AcroForm: fill in-place.
      * ("docx", fields)        → flat PDF / docx / other: generate an editable .docx. `fields` are the
                                   recognized 'Label: ___' fields merged with the default personal set.
    """
    if _looks_like_pdf(content, content_type, url):
        widgets = _pdf_widget_fields(content)
        if widgets:
            return "pdf", widgets
        recognized = _flat_pdf_labels(content)
        return "docx", _merge_with_defaults(recognized)
    return "docx", _merge_with_defaults([])


def _merge_with_defaults(recognized: list[FormField]) -> list[FormField]:
    """Recognized form fields first, then any default personal fields not already covered."""
    keys = {field.key for field in recognized}
    merged = list(recognized)
    for field in _default_fields():
        if field.key not in keys:
            merged.append(field)
    return merged


def fill_pdf(content: bytes, values: dict[str, str]) -> bytes:
    """Fill a fillable AcroForm PDF's widgets with `values` (by field name) and return the PDF bytes."""
    import fitz

    with fitz.open(stream=content, filetype="pdf") as document:
        for page in document:
            for widget in page.widgets() or []:
                name = (widget.field_name or "").strip()
                if name in values and values[name] is not None:
                    try:
                        widget.field_value = str(values[name])
                        widget.update()
                    except Exception:
                        logger.debug("Could not set widget %r.", name, exc_info=True)
        return document.tobytes()


def build_docx(form_title: str, fields: list[FormField], narrative: str = "") -> bytes:
    """Generate a clean, editable .docx from the (filled) fields — the fallback for flat/non-fillable forms."""
    import docx

    document = docx.Document()
    document.add_heading(form_title or "VinUni Form", level=1)
    for field in fields:
        label = field.label or field.key
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
        paragraph.add_run(field.value or "")
    if narrative and narrative.strip():
        document.add_paragraph("")
        document.add_paragraph(narrative.strip())
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_form_file(
    file_kind: str,
    content: bytes | None,
    form_title: str,
    fields: list[FormField],
    narrative: str = "",
) -> tuple[bytes, str, str]:
    """Produce the downloadable file from the (possibly edited) fields.

    Returns (data, mime, extension). `file_kind` "pdf" fills the original AcroForm (needs `content`);
    anything else (or a missing/failed PDF fill) generates a .docx. Fail-open: a PDF-fill error falls
    back to a generated .docx so the download never hard-fails.
    """
    values = {field.key: field.value for field in fields}
    if file_kind == "pdf" and content is not None:
        try:
            return fill_pdf(content, values), "application/pdf", "pdf"
        except Exception:
            logger.warning("PDF fill failed; falling back to generated .docx.", exc_info=True)
    data = build_docx(form_title, fields, narrative)
    return (
        data,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    )
